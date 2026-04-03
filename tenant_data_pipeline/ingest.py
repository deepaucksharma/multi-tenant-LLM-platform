"""
Ingest raw documents from tenant directories.
Supports: .txt, .pdf, .csv, .docx
Outputs cleaned text with metadata.
"""
import json
import hashlib
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import List, Optional
from datetime import datetime

from loguru import logger

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from tenant_data_pipeline.config import TENANTS, TenantConfig


@dataclass
class IngestedDocument:
    doc_id: str
    tenant_id: str
    source_file: str
    title: str
    topic: str
    content: str
    content_hash: str
    word_count: int
    char_count: int
    ingested_at: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def compute_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def extract_text_from_txt(filepath: Path) -> str:
    return filepath.read_text(encoding="utf-8", errors="replace")


def extract_text_from_pdf(filepath: Path) -> str:
    if PdfReader is None:
        logger.warning(f"pypdf not installed, skipping {filepath}")
        return ""
    reader = PdfReader(str(filepath))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def extract_text_from_csv(filepath: Path) -> str:
    if not PANDAS_AVAILABLE:
        logger.warning(f"pandas not installed, skipping {filepath}")
        return ""
    import pandas as pd
    df = pd.read_csv(filepath)
    rows = []
    for _, row in df.iterrows():
        row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
        rows.append(row_text)
    return "\n".join(rows)


def extract_text_from_docx(filepath: Path) -> str:
    if DocxDocument is None:
        logger.warning(f"python-docx not installed, skipping {filepath}")
        return ""
    doc = DocxDocument(str(filepath))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


EXTRACTORS = {
    ".txt": extract_text_from_txt,
    ".pdf": extract_text_from_pdf,
    ".csv": extract_text_from_csv,
    ".docx": extract_text_from_docx,
}


def parse_title_topic(text: str, filename: str) -> tuple:
    """Extract title and topic from document header if present."""
    title = filename
    topic = "general"
    lines = text.strip().split("\n")
    for line in lines[:5]:
        if line.startswith("TITLE:"):
            title = line.replace("TITLE:", "").strip()
        elif line.startswith("TOPIC:"):
            topic = line.replace("TOPIC:", "").strip()
    return title, topic


def ingest_tenant(tenant_id: str) -> List[IngestedDocument]:
    """Ingest all documents for a single tenant."""
    config = TENANTS[tenant_id]
    config.processed_dir.mkdir(parents=True, exist_ok=True)

    documents = []
    raw_dir = config.raw_dir

    if not raw_dir.exists():
        logger.warning(f"Raw directory not found: {raw_dir}")
        return documents

    for filepath in sorted(raw_dir.iterdir()):
        if filepath.name == "manifest.json":
            continue

        suffix = filepath.suffix.lower()
        extractor = EXTRACTORS.get(suffix)
        if extractor is None:
            logger.info(f"Skipping unsupported file: {filepath}")
            continue

        try:
            raw_text = extractor(filepath)
        except Exception as e:
            logger.error(f"Failed to extract {filepath}: {e}")
            continue

        if not raw_text.strip():
            logger.warning(f"Empty content from {filepath}")
            continue

        title, topic = parse_title_topic(raw_text, filepath.stem)
        content_hash = compute_hash(raw_text)
        doc_id = f"{tenant_id}_{content_hash}"

        doc = IngestedDocument(
            doc_id=doc_id,
            tenant_id=tenant_id,
            source_file=str(filepath.name),
            title=title,
            topic=topic,
            content=raw_text,
            content_hash=content_hash,
            word_count=len(raw_text.split()),
            char_count=len(raw_text),
            ingested_at=datetime.utcnow().isoformat(),
            file_type=suffix,
        )
        documents.append(doc)

    output_path = config.processed_dir / "ingested_documents.json"
    output_data = [asdict(d) for d in documents]
    output_path.write_text(json.dumps(output_data, indent=2), encoding="utf-8")
    logger.info(f"[{tenant_id}] Ingested {len(documents)} documents → {output_path}")

    return documents


def ingest_all_tenants() -> dict:
    """Ingest documents for all tenants."""
    results = {}
    for tenant_id in TENANTS:
        docs = ingest_tenant(tenant_id)
        results[tenant_id] = len(docs)
    return results


if __name__ == "__main__":
    results = ingest_all_tenants()
    print(f"Ingestion results: {results}")
